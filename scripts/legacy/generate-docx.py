#!/usr/bin/env python3
"""Minimal DOCX generator - no temp files"""
import sys
import os

# Try importing docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not available. Install with: pip install python-docx")
    sys.exit(1)

# Create and save directly without temp operations
try:
    doc = Document()

    # Title
    title = doc.add_heading("The AMI English Language Progression", level=1)
    doc.add_paragraph("From Zero to Independent Reading")
    doc.add_paragraph("For Tredoux")

    # Save directly
    outpath = "/sessions/keen-pensive-cerf/mnt/whale/AMI_English_Language_Progression.docx"
    doc.save(outpath)

    if os.path.exists(outpath):
        size = os.path.getsize(outpath)
        print(f"✓ Document created: {outpath} ({size} bytes)")
        sys.exit(0)
    else:
        print(f"ERROR: File not created")
        sys.exit(1)

except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
