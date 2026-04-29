#!/usr/bin/env python3
"""
Generate AMI English Language Progression Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
MONTESSORI_GREEN = RGBColor(46, 125, 50)
LIGHT_GREEN_HEX = "E8F5E9"
TEXT_DARK = RGBColor(27, 27, 27)
TEXT_LIGHT = RGBColor(85, 85, 85)

def shade_cell(cell, fill_color):
    """Add shading to a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading1(doc, text):
    """Add a Heading 1"""
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = MONTESSORI_GREEN
        run.font.size = Pt(28)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading2(doc, text):
    """Add a Heading 2"""
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MONTESSORI_GREEN
        run.font.size = Pt(24)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_text(doc, text, bold=False):
    """Add body text"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(doc, text):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    return p

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ========== TITLE PAGE ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("The AMI English Language Progression")
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = MONTESSORI_GREEN
title.paragraph_format.space_before = Pt(72)
title.paragraph_format.space_after = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("From Zero to Independent Reading")
subtitle_run.font.size = Pt(18)
subtitle_run.font.color.rgb = TEXT_LIGHT
subtitle_run.italic = True
subtitle.paragraph_format.space_after = Pt(48)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run("A Simple Guide to Understanding How Children Learn to Read in the Montessori Method")
desc_run.font.size = Pt(12)
desc_run.font.color.rgb = TEXT_LIGHT
desc.paragraph_format.space_before = Pt(48)
desc.paragraph_format.space_after = Pt(72)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run("For Tredoux")
author_run.font.size = Pt(12)
author_run.font.color.rgb = MONTESSORI_GREEN
author_run.italic = True

# Page break
doc.add_page_break()

# ========== SECTION 1: THE BIG IDEA ==========
add_heading1(doc, "The Big Idea: Why Montessori Does Everything Backwards (And It Makes Perfect Sense)")

add_body_text(doc, "If you've ever watched a child in a Montessori classroom, you've probably noticed something strange: the teacher is teaching them to WRITE before they READ. This seems totally backwards. How can a child write if they can't read?")

add_body_text(doc, "But here's the thing: it's not backwards at all. It's actually the most logical progression that exists. And once you understand why, everything about the Montessori language approach will make sense.")

add_heading2(doc, "Three Big Insights")

add_bullet(doc, "Writing is EXPRESSION. Reading is INTERPRETATION. Expressing comes before interpreting. A child naturally wants to MAKE something (write) before they want to UNDERSTAND something (read).")

add_bullet(doc, "Sounds matter more than letter names. We teach /mmm/ not 'em.' Because /mmm/ is what the child hears in words. Letter names are abstract and useless for sounding out words.")

add_bullet(doc, "The hand follows the brain. Before a child can write, their hand must be ready. But before their hand can be ready, their brain must understand what letters ARE. So we do: sounds → hand prep → writing → reading. In that order.")

add_body_text(doc, "One more thing that makes Montessori work: every material is self-correcting. A child knows immediately if they did it right. No confusion, no frustration, no waiting for the teacher to check their work.")

add_body_text(doc, "Ready? Let's watch a child named Marina grow from not knowing a single letter to reading her first book.")

doc.add_page_break()

# ========== SECTION 2: THE JOURNEY ==========
add_heading1(doc, "The Journey: Marina From Zero to Reader")

add_body_text(doc, "Marina is 3 years old. She doesn't know any letters. She can't hold a pencil properly. But she can listen. She can talk. And she's curious about the world. This is where we start.")

# Stage 1
add_heading2(doc, "Stage 1: Listening & Speaking (Ages 2.5–3.5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She listens. Really listens. To songs, poems, and conversations.")
add_bullet(doc, "She learns new words every day. The teacher names everything precisely.")
add_bullet(doc, "She plays with language—songs, nursery rhymes, finger plays.")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Her ears. That's it.")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "A child needs a RICH VOCABULARY before she can learn to read or write. She needs to know what words SOUND like. If you ask Marina to spell 'cat' without knowing the word 'cat' exists, she's lost. Reading and writing are not about inventing language. They're about ENCODING and DECODING language she already knows.")

add_body_text(doc, "The classroom is full of language. The teacher talks. Other children talk. Stories are read aloud. Songs are sung. Marina's ears are drinking it all in.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When Marina is spontaneously using new words, enjoying rhymes, and noticing sounds in words. When she laughs at silly word combinations. When she asks what things are called. She's ready.")

# Stage 2
add_heading2(doc, "Stage 2: Sound Games & I Spy (Ages 2.5–3.5, runs parallel with Stage 1)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She plays 'I Spy' games focused on SOUNDS, not letters.")
add_bullet(doc, "Teacher: 'I spy something beginning with /mmm/.'")
add_bullet(doc, "Marina guesses: mirror, muffin, milk.")
add_bullet(doc, "No letters shown. Just sounds.")

add_body_text(doc, "She learns to isolate sounds within words:", True)
add_bullet(doc, "Initial sounds: /sss/ in 'sun'")
add_bullet(doc, "Ending sounds: /t/ in 'cat'")
add_bullet(doc, "Middle sounds: /a/ in 'cat'")
add_bullet(doc, "How to blend sounds: /c/ + /a/ + /t/ = 'cat'")
add_bullet(doc, "How to segment words: 'cat' = /c/ + /a/ + /t/")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Her ears. Baskets of small objects (a marble, a fork, a button). Pictures in boxes.")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "Before Marina can read the word 'cat,' her brain needs to HEAR what /c/ + /a/ + /t/ SOUNDS like when you push them together. This is called PHONEMIC AWARENESS. Without it, letters are just random squiggles. With it, letters become a code for sounds she already knows.")

add_body_text(doc, "Marina doesn't need to see the letter 'c' yet. She just needs to know that /c/ is a sound that shows up in certain words.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When she can play I Spy confidently, find initial sounds easily, and can blend 2-3 sounds together into a word. When she's noticing sounds without prompting. She's ready for the tactile experience.")

# Stage 3
add_heading2(doc, "Stage 3: Metal Insets (Ages 3+, parallel with sound games)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She traces geometric shapes inside metal frames using a colored pencil.")
add_bullet(doc, "She stays inside the lines. She learns what LIGHTNESS OF TOUCH feels like.")
add_bullet(doc, "She draws straight lines, curves, circles, wavy lines.")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Ten metal frames with shapes inside. Colored pencils. Paper. Concentration.")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "You cannot write letters without the hand muscles, finger control, and lightness of touch that metal insets teach. Every stroke in every letter is made of curves and straight lines. Every stroke in a metal inset is practice for the strokes in letters.")

add_body_text(doc, "But here's the SECRET: metal insets are NOT about letter formation. They're about HAND PREPARATION. Marina doesn't know what letters are yet. She's just playing with shapes. But her hand is getting exactly the practice it needs.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When she can trace all ten frames confidently, stay inside the lines most of the time, and her hand is visibly steadier and lighter. When she's asking for more. She's ready to learn what letters are.")

# Stage 4
add_heading2(doc, "Stage 4: Sandpaper Letters (Ages 3–3.5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She feels a ROUGH sandpaper letter while the teacher says the SOUND.")
add_bullet(doc, "Teacher traces the letter with two fingers: 'This is /mmm/.'")
add_bullet(doc, "Marina traces it with two fingers, saying /mmm/ out loud.")
add_bullet(doc, "She traces it again. And again. Tactile + auditory + sight all at once.")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Sandpaper letters glued to wooden boards. Grouped in sets of 3–4 contrasting letters (not alphabetical order).")
add_bullet(doc, "Group 1: c, m, a, t (very different shapes and sounds)")
add_bullet(doc, "Group 2: s, r, i, p")
add_bullet(doc, "Group 3: b, f, o, g")
add_bullet(doc, "And so on.")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "Marina is meeting LETTERS for the first time. But not as abstract concepts. She FEELS them (rough), HEARS them (the sound), and SEES them (visually) all at the same time. Three pathways to the brain at once. This is POWERFUL.")

add_body_text(doc, "When her two fingers trace the sandpaper 'm,' the muscle memory of that movement goes into her brain. Later, when she picks up a pencil to WRITE an 'm,' her hand will remember. It will know which way to go, how much pressure to use, when to curve.")

add_body_text(doc, "Notice: we're teaching the SOUND of the letter (/mmm/), not its name ('em'). Because Marina needs to hear the sound when she's trying to decode words later.")

add_body_text(doc, "Notice: the letters aren't in alphabetical order. They're grouped by contrast. If you teach c, m, a, t together, Marina can't confuse c with m. They look totally different. They feel totally different. By the time you introduce letters that look similar (like b and d), Marina is already solid with the others.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When she can feel a letter, say its sound without the teacher showing her, and when she's asking to learn more. When she's spending time just tracing letters for fun, not because the teacher asked her to. She's ready to BUILD words.")

# Stage 5
add_heading2(doc, "Stage 5: Moveable Alphabet (Ages 3.5–4) — THE BREAKTHROUGH")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She sees wooden tiles with letters on them.")
add_bullet(doc, "Teacher says a word: 'cat.'")
add_bullet(doc, "Marina finds the 'c' tile, the 'a' tile, the 't' tile.")
add_bullet(doc, "She arranges them left to right: c-a-t.")
add_bullet(doc, "Teacher reads what she built: 'cat.'")
add_bullet(doc, "Marina just READ her own writing. Without ever picking up a pencil.")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "The Moveable Alphabet — hundreds of small wooden tiles with letters, organized by sound in a wooden box.")

add_body_text(doc, "Why this step is THE BREAKTHROUGH:", True)
add_body_text(doc, "This is the moment everything changes. Marina has just discovered that letters are a CODE. She can BUILD a word with tiles. The word appears in front of her. She reads it.")

add_body_text(doc, "She's WRITING (encoding). But she's doing it without needing perfect handwriting. Her brain is learning the SEQUENCE of letters in words. Later, her hand will put a pencil to paper and write those sequences. But first, her brain needs to GET IT.")

add_body_text(doc, "The moveable alphabet is genius because:")
add_bullet(doc, "It's FAST. Marina doesn't struggle with pencil control. The letters just slide into place.")
add_bullet(doc, "It's OBVIOUS when you make a mistake. The word doesn't sound right. Marina corrects it herself.")
add_bullet(doc, "It's DEEP LEARNING. Marina is working with REAL words she already knows how to SAY.")

add_body_text(doc, "From this point on, Marina understands: LETTERS = SOUNDS. SEQUENCES OF LETTERS = WORDS.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When she can build 5+ CVC words confidently, when she's enjoying it, when she's asking to build more words. When she understands that the tiles REPRESENT sounds. She's ready to READ those same words in print.")

# Stage 6
add_heading2(doc, "Stage 6: Pink Series — CVC Reading (Ages 4–4.5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She looks at a card with a 3-letter word printed on it: 'cat.'")
add_bullet(doc, "She reads it. She already KNOWS this word because she BUILT it with moveable alphabet.")
add_bullet(doc, "She matches the card to a picture of a cat.")
add_bullet(doc, "She moves to a word list with 10 cat-words: cat, can, car, cap, cam, cad, cot, cub, cut, cap.")
add_bullet(doc, "She reads them. All CVC (consonant-vowel-consonant). All phonetic. All decodable.")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Pink Series books. Three components:")
add_bullet(doc, "Object boxes — real small objects (a bat, a mat, a rat, a cat toy)")
add_bullet(doc, "Picture cards — images of those objects")
add_bullet(doc, "Word cards — the printed words")
add_bullet(doc, "Word lists — pages with 10 related words")
add_bullet(doc, "Booklets — simple sentences and short stories using Pink words")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "Marina has BUILT these words. Now she READS them in print. Reading her own writing. Then reading other people's writing. The progression is gentle and familiar.")

add_body_text(doc, "Pink words are all PHONETIC. Every sound is decodable. 'cat' = /c/ + /a/ + /t/. Marina's brain makes the connection instantly because she's already BUILT this word.")

add_body_text(doc, "The progression goes:")
add_bullet(doc, "Objects (most concrete) → Pictures (more abstract) → Words (most abstract) → Sentences → Stories")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When she's confidently reading all the Pink Series booklets, when she's not sounding out every letter anymore, when reading feels NORMAL to her. When she starts reading words she HASN'T been taught. That's when you know her brain has cracked the code. She's ready for words that are more complex.")

# Stage 7
add_heading2(doc, "Stage 7: Blue Series — Blends (Ages 4–5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She reads words with consonant BLENDS: frog, stop, hand, jump, trip, skip.")
add_bullet(doc, "Two consonants together. But each consonant KEEPS ITS SOUND.")
add_bullet(doc, "She reads: /f/ + /r/ + /o/ + /g/ = 'frog.'")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Blue Series books. Same structure as Pink:")
add_bullet(doc, "Object boxes → Pictures → Word cards → Word lists → Booklets")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "Marina is ready for more complex phonetic words. But we're not introducing new SOUNDS yet. Just new COMBINATIONS of sounds she already knows. Two consonants that keep their individual sounds.")

add_body_text(doc, "This is still 100% phonetic and decodable. Marina's brain can handle the complexity now because she's mastered simple CVC patterns.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When Blue Series feels easy. When she's not struggling with the blends. When she's confidently reading Blue booklets. She's ready to learn about NEW sounds.")

# Stage 8
add_heading2(doc, "Stage 8: Green Series — Phonograms/Digraphs (Ages 4.5–5.5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She reads words with DIGRAPHS — letter combinations that make NEW sounds:")
add_bullet(doc, "'sh' makes /sh/ (not /s/ + /h/)")
add_bullet(doc, "'ch' makes /ch/ (not /c/ + /h/)")
add_bullet(doc, "'th' makes /th/ (not /t/ + /h/)")
add_bullet(doc, "'ai' makes /ai/ (not /a/ + /i/)")
add_bullet(doc, "'oa' makes /o/ (not /o/ + /a/)")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Green Series books. Same structure as Pink and Blue.")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "Marina has mastered SINGLE SOUNDS (Pink) and SOUND COMBINATIONS where each sound stays individual (Blue). Now she's learning about SOUND PUZZLES — two letters that make a COMPLETELY DIFFERENT sound than you'd expect.")

add_body_text(doc, "This is still phonetic (once you know the 'rule' of the digraph). But it requires Marina to understand that some letter pairs are SPECIAL. They're puzzle pieces.")

add_body_text(doc, "How you know she's ready for the next step:", True)
add_body_text(doc, "When Green Series feels natural. When she's automatically recognizing digraphs in words. When she's reading Green booklets fluently. She's ready for exceptions to the rules.")

# Stage 9
add_heading2(doc, "Stage 9: Puzzle Words / Sight Words (Introduced after Pink confidence)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She learns words that BREAK THE RULES.")
add_bullet(doc, "'the' — she can't decode this. It's a puzzle.")
add_bullet(doc, "'said' — she can't decode this. It's a puzzle.")
add_bullet(doc, "'was' — she can't decode this. It's a puzzle.")
add_bullet(doc, "She memorizes these words because they appear SO OFTEN in books.")

add_body_text(doc, "Why puzzle words come AFTER Pink (not before):", True)
add_body_text(doc, "Marina needs a PHONETIC FOUNDATION first. She needs to understand the basic code. Only then can you introduce exceptions. If you start with puzzle words, she has no foundation. No understanding. Just confusion.")

add_body_text(doc, "In Montessori, we do rules before exceptions. ALWAYS.")

add_body_text(doc, "How you know she's ready:", True)
add_body_text(doc, "When she's fluent with Pink. When she starts encountering puzzle words in books and asks how to read them. That's the signal: introduce them now.")

# Stage 10
add_heading2(doc, "Stage 10: Handwriting on Paper (Ages 4.5–5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She picks up a pencil.")
add_bullet(doc, "She writes letters on paper. Not because the teacher forced her, but because she's READY.")
add_bullet(doc, "Her hand is steady (metal insets). Her brain knows the letters (sandpaper). She's been 'writing' with tiles (moveable alphabet). Writing with a pencil comes naturally.")

add_body_text(doc, "What materials she uses:", True)
add_bullet(doc, "Paper, pencil, lined worksheets.")

add_body_text(doc, "Why this comes LAST, not first:", True)
add_body_text(doc, "The pencil is the HARDEST tool. It requires the steadiest hand, the most muscle control, the most endurance. So we prepare first (metal insets), then teach (sandpaper), then practice without the pencil (moveable alphabet), and ONLY THEN do we ask her to write with a pencil.")

add_body_text(doc, "By the time Marina holds a pencil, she's already comfortable with letters. She's already confident. The pencil is just the final step.")

add_body_text(doc, "How you know she's ready:", True)
add_body_text(doc, "When she's asking to write. When her hand is visibly steady and controlled. When she's EAGER, not reluctant. That's when you introduce pencil writing. Not before.")

# Stage 11
add_heading2(doc, "Stage 11: Command Cards (Ages 4.5–5)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She reads a card that says: 'Hop to the door.'")
add_bullet(doc, "She READS the words.")
add_bullet(doc, "She UNDERSTANDS the meaning.")
add_bullet(doc, "She DOES the action.")

add_body_text(doc, "Why this step exists:", True)
add_body_text(doc, "Marina has been reading LISTS and PICTURES. Now she's reading for MEANING. She's reading for COMPREHENSION. The proof is in the action. If she understood, she can do the command.")

add_body_text(doc, "This is where reading becomes USEFUL instead of just PRACTICE.")

add_body_text(doc, "How you know she's ready:", True)
add_body_text(doc, "When she's fluent with Pink and Blue series. When she's not sounding out every word anymore. She's ready to read for meaning.")

# Stage 12
add_heading2(doc, "Stage 12: Grammar (Ages 5–6)")

add_body_text(doc, "What Marina does:", True)
add_bullet(doc, "She learns parts of speech using PHYSICAL SYMBOLS:")
add_bullet(doc, "A black triangle = NOUN (a person, place, or thing)")
add_bullet(doc, "A red circle = VERB (an action)")
add_bullet(doc, "A blue ball = ADJECTIVE (a description)")
add_bullet(doc, "She reads sentences and places symbols above each word to show what kind of word it is.")

add_body_text(doc, "Why grammar comes LAST:", True)
add_body_text(doc, "Marina has been reading for months. Her brain KNOWS how words work. Now she's learning to THINK ABOUT how words work. Grammar is the conscious understanding of something she's already doing unconsciously.")

add_body_text(doc, "And notice: it's not abstract. It's PHYSICAL. She can TOUCH the symbols. She can MOVE them around. Grammar is not something that happens in a textbook. It's something she can DO.")

doc.add_page_break()

# ========== SECTION 3: PINK/BLUE/GREEN ==========
add_heading1(doc, "The Pink, Blue, and Green Systems: A Visual Summary")

add_body_text(doc, "The foundation of reading in Montessori is the THREE SERIES of books. Each series introduces a new layer of complexity. Each series follows the same progression.")

# Table
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = "Series"
header_cells[1].text = "What It Teaches"
header_cells[2].text = "Example Words"
header_cells[3].text = "Age"

for cell in header_cells:
    shade_cell(cell, MONTESSORI_GREEN)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
data = [
    ("PINK", "CVC words (3-letter, phonetic)", "cat, dog, pin, hen, bat", "4–4.5 y"),
    ("BLUE", "Blends (consonant clusters)", "frog, stop, hand, jump", "4–5 y"),
    ("GREEN", "Digraphs (new sounds)", "fish, boat, chair, rain", "4.5–5.5 y"),
]

for i, (series, teaches, examples, age) in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = series
    row_cells[1].text = teaches
    row_cells[2].text = examples
    row_cells[3].text = age

    if i % 2 == 1:
        shade_cell(row_cells[0], LIGHT_GREEN_HEX)

add_body_text(doc, "Each series follows the exact same progression:")
add_bullet(doc, "OBJECTS — Real, touchable things (most concrete)")
add_bullet(doc, "PICTURES — Images of those objects (more abstract)")
add_bullet(doc, "WORD CARDS — Single words written out (even more abstract)")
add_bullet(doc, "PHRASES & SENTENCES — Words combined into meaning (most abstract)")
add_bullet(doc, "BOOKLETS — Stories using the words from the series")

add_body_text(doc, "This progression from CONCRETE to ABSTRACT is the Montessori way. First, touch something real. Then, see a picture of it. Then, see the WORD that names it. By the time the child reaches the word, they already KNOW what it means.")

add_body_text(doc, "Each series also gets HARDER. But children move at their own pace. Some children race through Pink in a few months. Others take longer. This is NORMAL and FINE. By age 5.5, most children have worked through all three series and are reading simple chapter books independently.")

doc.add_page_break()

# ========== SECTION 4: SANDPAPER LETTER GROUPS ==========
add_heading1(doc, "The Sandpaper Letter Groups: The Precise Order")

add_body_text(doc, "The order in which letters are introduced matters. They're NOT in alphabetical order. They're grouped by contrast and frequency.")

add_body_text(doc, "Letters that look similar are taught FAR APART. If you teach 'b' and 'd' at the same time, Marina will confuse them for months. But if you teach 'b' with very different-looking letters, she'll remember it immediately.")

add_body_text(doc, "Letters that are common are taught first. Letters that are rare are taught last.")

add_heading2(doc, "Group 1: c, m, a, t")
add_body_text(doc, "The first sounds. Very different from each other. All used in common CVC words (cat, mat, can, cat, act).")

add_heading2(doc, "Group 2: s, r, i, p")
add_body_text(doc, "Still very contrasting. The letter 'i' is the second vowel introduced. 's' is a common initial sound. 'r' and 'p' are distinct consonants.")

add_heading2(doc, "Group 3: b, f, o, g")
add_body_text(doc, "More consonants, another vowel. 'b' is now introduced far from 'd' (which comes later). 'o' is the third vowel.")

add_heading2(doc, "Group 4: h, j, u, l")
add_body_text(doc, "'u' is the fourth vowel. 'h' is common. 'j' and 'l' are distinct.")

add_heading2(doc, "Group 5: d, w, e, n")
add_body_text(doc, "'d' is now introduced (far from 'b'). 'e' is the fifth vowel. 'w' and 'n' are common consonants.")

add_heading2(doc, "Group 6: k, q, v, x, y, z")
add_body_text(doc, "The least common letters. Introduced after children are solid with all the others. 'y' because it's sometimes a vowel. 'q' because it almost always comes with 'u.' 'x,' 'v,' and 'z' because they're rare and easy to confuse once you know other letters.")

add_body_text(doc, "Why this matters:", True)
add_body_text(doc, "You follow this order RELIGIOUSLY. It's not a suggestion. It's been proven to work. When you introduce letters in contrast, children learn faster and make fewer mistakes.")

doc.add_page_break()

# ========== SECTION 5: WHY THIS ORDER WORKS ==========
add_heading1(doc, "Why This Order Works: The Deep Logic")

add_heading2(doc, "1. Sound Before Symbol")
add_body_text(doc, "We teach /mmm/ before we teach the letter 'm.' We teach EARS before EYES. Because ears are the sense that matters for reading. Reading is decoding sounds. If you jump straight to the letter, the child doesn't understand why the letter matters.")

add_heading2(doc, "2. Encoding Before Decoding")
add_body_text(doc, "We teach WRITING before READING. But we start with the moveable alphabet (writing without a pencil) before we ask children to pick up a pencil. Because encoding (putting thoughts into code) is easier than decoding (breaking code into thoughts). A child wants to MAKE before they want to UNDERSTAND.")

add_heading2(doc, "3. Concrete Before Abstract")
add_body_text(doc, "We start with real OBJECTS (a cat), then PICTURES of objects (a drawing of a cat), then the WORD (the letters c-a-t). Because the concrete is something the child already knows. The abstract (the word) connects to something real.")

add_heading2(doc, "4. Motor Prep Runs in Parallel")
add_body_text(doc, "While Marina is learning sounds, her hands are getting ready (metal insets). By the time she learns letters, her hand is READY to trace them. By the time she learns to write with a pencil, her hand is STRONG enough. Everything is timed perfectly.")

add_heading2(doc, "5. Phonetic First, Exceptions Second")
add_body_text(doc, "We teach the RULES before the EXCEPTIONS. Pink series (100% phonetic) before puzzle words. Blue series (still all phonetic blends) before trigraphs that have weird sounds. If you try to teach exceptions first, children have no foundation. With a phonetic foundation, exceptions make sense.")

add_heading2(doc, "6. Self-Correction Built Into Everything")
add_body_text(doc, "The moveable alphabet tiles make a word that either SOUNDS right or doesn't. The sandpaper letters are rough — you can FEEL if you traced them right. The picture cards either MATCH the word or they don't. Children know immediately if they did it right. No waiting for the teacher. No frustration.")

add_heading2(doc, "7. Nothing Is Wasted")
add_body_text(doc, "Metal insets are NOT 'finger painting for fine motor skills.' They're TRAINING your hand for the exact strokes you'll use in letters. The sound games are NOT 'fun activities.' They're BUILDING the phonemic awareness you need to decode words. Everything has a purpose. Everything connects.")

add_heading2(doc, "8. Pacing Is Individual")
add_body_text(doc, "Some children zoom through Pink in two months. Some take six months. Both are NORMAL. We never rush. We never slow down artificially. We follow the CHILD. When Marina is ready to move to Blue, we move. When she needs more time with Pink, she gets it. No timelines. No frustration.")

# Conclusion
add_body_text(doc, "")
add_body_text(doc, "This is the Montessori logic. Simple. Logical. Built on how children's brains actually work. Not how we THINK they should work. Not what's fashionable. What WORKS.")

add_body_text(doc, "And the proof is in the outcome: children who follow this progression rarely struggle with reading. By age 5 or 6, they're reading independently. Not because they were forced. Because they were ready. Because everything was prepared. Because their brain, their hand, and their spirit all lined up at the right moment.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("—")
run.font.size = Pt(14)
run.font.color.rgb = MONTESSORI_GREEN
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(10)

add_body_text(doc, "That's Marina's journey. From silence to reading. Simple. Logical. Beautiful.")

# Save document
doc.save("/sessions/keen-pensive-cerf/mnt/whale/AMI_English_Language_Progression.docx")
print("✓ Document created successfully!")
print("Saved to: /sessions/keen-pensive-cerf/mnt/whale/AMI_English_Language_Progression.docx")
